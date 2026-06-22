"""
Módulo de análise automática de editais em PDF.
Extrai campos via regex e gera relatório Word (.docx).
"""

import io
import re
from collections import Counter
from datetime import datetime

import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Extração de texto ──────────────────────────────────────────────────────────

def extrair_texto_pdf(arquivo) -> str:
    """Recebe file-like object, retorna texto completo do PDF."""
    paginas = []
    with pdfplumber.open(arquivo) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                paginas.append(t)
    return "\n".join(paginas)


# ── Extração de campos via regex ───────────────────────────────────────────────

def extrair_campos(texto: str) -> dict:
    campos = {}

    # Número do edital
    for pat in [
        r'[Ee]dital\s+n[º°\.°]\s*[\d./\-]+(?:/\d+)?',
        r'EDITAL\s+N[Oº°]\s*[\d./\-]+',
        r'[Pp]regão\s+[Ee]letrônico\s+n[º°\.]\s*[\d./\-]+',
        r'[Cc]oncorrência\s+[Ee]letrônica?\s+n[º°\.]\s*[\d./\-]+',
        r'[Pp]rocesso\s+(?:[Aa]dministrativo\s+)?n[º°\.]\s*[\d./\-\s]+',
    ]:
        m = re.search(pat, texto)
        if m:
            campos["numero_edital"] = m.group(0).strip()
            break

    # CNPJ
    cnpjs = re.findall(r'\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}', texto)
    if cnpjs:
        campos["cnpj"] = cnpjs[0]

    # Órgão licitante
    for pat in [
        r'(?:UASG|Uasg)[:\s]+\d+\s*[-–]\s*([^\n]{5,80})',
        r'(?:Órgão|ÓRGÃO)\s*(?:Licitante|licitante)?[:\s]+([^\n]{5,80})',
        r'(?:Contratante|CONTRATANTE)[:\s]+([^\n]{5,80})',
    ]:
        m = re.search(pat, texto)
        if m:
            campos["orgao"] = m.group(1).strip()
            break

    # Valor estimado
    for pat in [
        r'[Vv]alor\s+[Tt]otal\s+[Ee]stimado[:\s]*R\$\s*([\d.,]+)',
        r'[Vv]alor\s+[Gg]lobal\s+[Ee]stimado[:\s]*R\$\s*([\d.,]+)',
        r'[Vv]alor\s+[Ee]stimado[:\s]*R\$\s*([\d.,]+)',
        r'[Pp]reço\s+[Gg]lobal\s+[Ee]stimado[:\s]*R\$\s*([\d.,]+)',
        r'[Oo]rçamento\s+[Ee]stimado[:\s]*R\$\s*([\d.,]+)',
    ]:
        m = re.search(pat, texto)
        if m:
            campos["valor_estimado"] = f"R$ {m.group(1)}"
            break

    # Data de abertura
    for pat in [
        r'[Dd]ata\s+(?:e\s+[Hh]ora\s+)?(?:de\s+)?[Aa]bertura[:\s]+(\d{2}/\d{2}/\d{4})',
        r'[Aa]bertura\s+(?:das\s+[Pp]ropostas?)?[:\s]+(\d{2}/\d{2}/\d{4})',
        r'[Ss]essão\s+[Pp]ública[:\s]+(\d{2}/\d{2}/\d{4})',
    ]:
        m = re.search(pat, texto)
        if m:
            campos["data_abertura"] = m.group(1)
            break

    # Data de encerramento
    for pat in [
        r'[Ee]ncerramento\s+(?:das?\s+[Pp]ropostas?)?[:\s]+(\d{2}/\d{2}/\d{4})',
        r'[Ll]imite\s+(?:para\s+)?[Ee]nvio[:\s]+(\d{2}/\d{2}/\d{4})',
        r'[Pp]razo\s+(?:final\s+)?[Pp]ara\s+[Ee]nvio[:\s]+(\d{2}/\d{2}/\d{4})',
    ]:
        m = re.search(pat, texto)
        if m:
            campos["data_encerramento"] = m.group(1)
            break

    # Prazo de execução
    for pat in [
        r'[Pp]razo\s+(?:de\s+)?[Ee]xecu[çc][ãa]o[:\s]+([\d]+\s*(?:dias?|meses?|anos?)(?:\s+(?:corridos?|úteis?))?)',
        r'[Pp]razo\s+[Cc]ontratual[:\s]+([\d]+\s*(?:dias?|meses?|anos?))',
        r'(\d+\s*(?:dias?\s+corridos?|dias?\s+úteis?|meses?\s+corridos?))',
    ]:
        m = re.search(pat, texto, re.IGNORECASE)
        if m:
            campos["prazo_execucao"] = m.group(1).strip()
            break

    # Modalidade
    for modalidade in [
        "Concorrência Eletrônica", "Concorrência Presencial", "Concorrência",
        "Pregão Eletrônico", "Pregão Presencial", "Pregão",
        "RDC Eletrônico", "RDC Presencial", "RDCI", "RDC",
        "Diálogo Competitivo", "Dispensa de Licitação",
    ]:
        if re.search(re.escape(modalidade), texto, re.IGNORECASE):
            campos["modalidade"] = modalidade
            break

    # Critério de julgamento
    for criterio in ["Menor Preço", "Técnica e Preço", "Melhor Técnica", "Maior Desconto"]:
        if re.search(re.escape(criterio), texto, re.IGNORECASE):
            campos["criterio_julgamento"] = criterio
            break

    # Garantia contratual
    for pat in [
        r'[Gg]arantia\s+(?:[Cc]ontratual|de\s+[Ee]xecu[çc][ãa]o)[:\s]+([\d,]+\s*%)',
        r'([\d,]+\s*%)\s*(?:do\s+)?(?:valor\s+)?(?:do\s+)?[Cc]ontrato',
    ]:
        m = re.search(pat, texto, re.IGNORECASE)
        if m:
            campos["garantia"] = m.group(1)
            break

    # Objeto
    for pat in [
        r'[Oo]bjeto\s+da\s+[Ll]icita[çc][ãa]o[:\s]+([^\n]{20,600})',
        r'[Oo]bjeto\s+do\s+[Cc]ontrato[:\s]+([^\n]{20,600})',
        r'[Oo]bjeto[:\s]+([^\n]{20,600})',
    ]:
        m = re.search(pat, texto)
        if m:
            campos["objeto"] = m.group(1).strip()[:500]
            break

    # UF mais frequente no documento
    ufs = re.findall(
        r'\b(AC|AL|AP|AM|BA|CE|DF|ES|GO|MA|MT|MS|MG|PA|PB|PR|PE|PI|RJ|RN|RS|RO|RR|SC|SP|SE|TO)\b',
        texto,
    )
    if ufs:
        campos["uf"] = Counter(ufs).most_common(1)[0][0]

    # Reajuste
    for pat in [
        r'[Rr]eajuste[:\s]+([^\n]{10,120})',
        r'[Íí]ndice\s+de\s+[Rr]eajuste[:\s]+([^\n]{10,80})',
        r'(INCC[^\n]{0,60})',
        r'(IPCA[^\n]{0,60})',
        r'(IGPM[^\n]{0,60})',
    ]:
        m = re.search(pat, texto)
        if m:
            campos["reajuste"] = m.group(1).strip()[:120]
            break

    # Qualificação técnica (atestados)
    trecho_qt = ""
    for pat in [
        r'[Qq]ualifica[çc][ãa]o\s+[Tt]écnica.{0,1500}',
        r'[Aa]testado.{0,800}',
    ]:
        m = re.search(pat, texto, re.DOTALL)
        if m:
            trecho_qt = m.group(0)[:600].strip()
            break
    if trecho_qt:
        campos["qualificacao_tecnica"] = trecho_qt

    return campos


# ── Geração do relatório Word ──────────────────────────────────────────────────

AZUL     = RGBColor(0x1F, 0x4E, 0x78)
CINZA    = RGBColor(0x70, 0x70, 0x70)
LARANJA  = RGBColor(0xC0, 0x50, 0x20)


def _add_heading(doc, text, level=1, color=None):
    color = color or AZUL
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def _add_table_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    vtext = value if value else "Não identificado automaticamente"
    row.cells[1].text = vtext

    r0 = row.cells[0].paragraphs[0].runs[0]
    r0.font.bold = True
    r0.font.size = Pt(10)

    r1 = row.cells[1].paragraphs[0].runs[0]
    r1.font.size = Pt(10)
    if not value:
        r1.font.color.rgb = CINZA


def gerar_relatorio_docx(campos: dict, nome_arquivo: str = "") -> io.BytesIO:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Cabeçalho
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("S.A. PAULISTA — ANÁLISE DE EDITAL")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = CINZA

    if nome_arquivo:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f"Arquivo: {nome_arquivo}")
        run3.font.size = Pt(9)
        run3.font.color.rgb = CINZA

    doc.add_paragraph()

    # 1. Identificação
    _add_heading(doc, "1. IDENTIFICAÇÃO DA LICITAÇÃO")
    t1 = doc.add_table(rows=0, cols=2)
    t1.style = "Table Grid"
    t1.columns[0].width = Cm(6)
    t1.columns[1].width = Cm(11)

    ident = [
        ("Número do Edital",         campos.get("numero_edital", "")),
        ("Órgão Licitante",          campos.get("orgao", "")),
        ("CNPJ",                     campos.get("cnpj", "")),
        ("UF",                       campos.get("uf", "")),
        ("Modalidade",               campos.get("modalidade", "")),
        ("Critério de Julgamento",   campos.get("criterio_julgamento", "")),
    ]
    for label, value in ident:
        _add_table_row(t1, label, value)

    doc.add_paragraph()

    # 2. Objeto
    _add_heading(doc, "2. OBJETO DO CONTRATO")
    objeto = campos.get("objeto", "")
    p_obj = doc.add_paragraph(objeto if objeto else "Não identificado automaticamente.")
    if not objeto:
        p_obj.runs[0].font.color.rgb = CINZA

    doc.add_paragraph()

    # 3. Financeiro e Prazos
    _add_heading(doc, "3. DADOS FINANCEIROS E PRAZOS")
    t2 = doc.add_table(rows=0, cols=2)
    t2.style = "Table Grid"
    t2.columns[0].width = Cm(6)
    t2.columns[1].width = Cm(11)

    fin = [
        ("Valor Estimado Total",  campos.get("valor_estimado", "")),
        ("Prazo de Execução",     campos.get("prazo_execucao", "")),
        ("Garantia Contratual",   campos.get("garantia", "")),
        ("Reajuste",              campos.get("reajuste", "")),
        ("Data de Abertura",      campos.get("data_abertura", "")),
        ("Data de Encerramento",  campos.get("data_encerramento", "")),
    ]
    for label, value in fin:
        _add_table_row(t2, label, value)

    doc.add_paragraph()

    # 4. Qualificação técnica
    qt = campos.get("qualificacao_tecnica", "")
    if qt:
        _add_heading(doc, "4. QUALIFICAÇÃO TÉCNICA (ATESTADOS)")
        p_qt = doc.add_paragraph(qt)
        p_qt.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # 5. Campos não identificados
    todos = ident + fin
    nao_id = [label for label, val in todos if not val]
    if nao_id:
        _add_heading(doc, "5. CAMPOS NÃO IDENTIFICADOS AUTOMATICAMENTE", level=2, color=LARANJA)
        p_aviso = doc.add_paragraph(
            "Os campos abaixo não foram localizados. Verificar manualmente no PDF original:"
        )
        p_aviso.runs[0].font.size = Pt(10)
        for campo in nao_id:
            doc.add_paragraph(f"• {campo}", style="List Bullet")

    # Nota de rodapé
    doc.add_paragraph()
    p_nota = doc.add_paragraph()
    run_nota = p_nota.add_run(
        "Nota: Relatório gerado por extração automática de texto (regex). "
        "Confirme sempre os dados no edital original antes de qualquer decisão."
    )
    run_nota.font.size = Pt(8)
    run_nota.font.color.rgb = CINZA
    run_nota.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
